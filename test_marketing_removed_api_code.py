import openai
import docx
import random
import time

# Set your OpenAI API key here
openai.api_key = 'redacted'
# List of topics
topics = ["marketing", "branding", "digital marketing"]

# Generate article using GPT-3
def generate_article(topic):
    prompt = f"Write an article about {topic}."
    response = openai.Completion.create(
        engine="text-davinci-003",  # You can choose the engine that suits your needs
        prompt=prompt,
        max_tokens=500  # Adjust this as needed
    )
    return response.choices[0].text.strip()

# Save article to a Word document
def save_to_word(article, topic, article_number):
    doc = docx.Document()
    doc.add_heading(f"{topic.capitalize()} Article {article_number}", level=1)
    doc.add_paragraph(article)
    doc.save(f"Documents/{topic}_article_{article_number}.docx")

# Generate and save 3 articles
for article_number in range(1, 4):
    random_topic = random.choice(topics)
    article = generate_article(random_topic)
    save_to_word(article, random_topic, article_number)
    print(f"Article {article_number} on {random_topic} generated and saved.")

    # Pause for a day
    time.sleep(24 * 60 * 60)  # Sleep for 24 hours
